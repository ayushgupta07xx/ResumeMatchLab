"""DOCX text extraction via python-docx (tables included)."""

from __future__ import annotations

import io


def parse_docx(data: bytes) -> tuple[str, str]:
    """Return ``(text, parser_used)`` from a .docx byte stream."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip(), "python-docx"
